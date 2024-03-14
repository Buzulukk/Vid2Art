"""
this project was developed during the period from June 23rd to June 25th, 2023, as part of the 'Profbuh hackathon'

i cannot be held responsible for the further functionality of this code, as it uses many third-party libraries

also you should put orig_vid.mp4 file of video you wanted to process in the directory with that .py file
(just the libraries for downloading videos refused to work during the hackathon, lol)
"""

"""
if you want to run the code, you first need to
enter the following lines into the console:

pip install youtube-transcript-api
pip install openai
pip install python-docx
pip install opencv-python
pip install telebot
pip install requests
pip install aiogram
pip install python-telegram-bot
pip install pyTelegramBotAPI

this will install the required libraries
"""


from youtube_transcript_api import YouTubeTranscriptApi  # lib for subs
import openai                                            # lib for chatGPT
import telebot, requests, re                             # lib for telegram
from telegram.ext import Updater, CommandHandler, MessageHandler, filters, CallbackContext
from aiogram.utils import *
import telegram
import docx                                              # lib for .docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

import math
import os


# Api key Telegram and OpenAI
bot = telebot.TeleBot('')
openai.api_key = ""

# welcome message
@bot.message_handler(commands=['start'])
def start_message(message):
    bot.send_message(message.chat.id, "Hello! Please enter the link to the YouTube video:")


# checking if the user entered a link
@bot.message_handler(func=lambda message: True)
def check_link_message(linkk):
    links = linkk.text
    if check_links(links) & (("youtube.com/" in str(links)) | ("youtu.be/" in str(links))):
        bot.send_message(linkk.chat.id, "Excellent choice! Please wait a moment, your request is being processed. This may take a few minutes.")
        link = str(links)

        def cut_link(link):  # the link truncation function
            return link[len(link) - 11:]

        if len(link) > 11:  # if the link is full we truncate it
            link = cut_link(link)

        URL_base = "https://youtu.be/" + link + "?t="  # used for creating hyperlinks

        video_path = "orig_vid.mp4"  # the name of the downloaded video for screenshots

        def capture_screenshot(video_path, timecode, i):  # function for screenshots
            current_dir = os.path.dirname(os.path.abspath(__file__))

            video = cv2.VideoCapture(video_path)

            # set timecode for screenshots
            video.set(cv2.CAP_PROP_POS_MSEC, timecode * 1000)
            # take a screenshot
            success, image = video.read()

            screenshot_path = os.path.join(current_dir, ("screenshot_" + link + "_" + str(i) + ".jpg"))
            cv2.imwrite(screenshot_path, image)
            
            video.release()

        def add_hyperlink(paragraph, url, text, color, underline):  # function for hyperlink

            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')

            # set a color
            if not color is None:
                c = docx.oxml.shared.OxmlElement('w:color')
                c.set(docx.oxml.shared.qn('w:val'), color)
                rPr.append(c)

            # underline
            if not underline:
                u = docx.oxml.shared.OxmlElement('w:u')
                u.set(docx.oxml.shared.qn('w:val'), 'none')
                rPr.append(u)

            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)

            paragraph._p.append(hyperlink)

            return hyperlink

        def create_part(subs_txt, begining_of_part, num_of_part):

            # title generation
            messages = [  # the instruction for chatGPT without its response
                {"role": "system",
                 "content": "Озаглавь текст в несколько слов"}
            ]

            messages.append({"role": "user", "content": subs_txt})

            completion = openai.ChatCompletion.create(  # generation of answer
                model="gpt-3.5-turbo-16k",
                messages=messages,
                n=1,
                temperature=0
            )

            chat_response_title = completion.choices[0].message.content

            # main text generation
            messages = [  # the instruction for chatGPT without its response
                {"role": "system",
                 "content": "Перескажи текст в одном абзаце"}
            ]

            messages.append({"role": "user", "content": subs_txt})  

            completion = openai.ChatCompletion.create(  # generation of answer
                model="gpt-3.5-turbo-16k",
                messages=messages,
                n=1,
                temperature=0
            )

            chat_response = completion.choices[0].message.content

            # adding a timecode, title, and text to a .docx file
            new_paragraph = result_docx.add_paragraph()
            _ = add_hyperlink(new_paragraph, URL_base + str(begining_of_part),
                              ('[ ' + str(begining_of_part // 60) + ':' + str(begining_of_part % 60) + ' ]'), "Blue",
                              True)
            title_of_part = result_docx.add_paragraph()
            runner = title_of_part.add_run(chat_response_title)
            runner.bold = True
            result_docx.add_paragraph(chat_response + '\n\n')

        language = ['en']

        # if it finds russian among the subtitles then the article will be in russian, otherwise English
        transcript_list = YouTubeTranscriptApi.list_transcripts(link)  
        for transcript in transcript_list:  
            if transcript.language_code == 'ru':
                language = ['ru']
                break

        result_docx = docx.Document()

        # adding the point "Based on:"
        result_docx.add_paragraph("Основано на:")
        new_paragraph = result_docx.add_paragraph()
        _ = add_hyperlink(new_paragraph, "https://www.youtube.com/watch?v=" + link,
                          "https://www.youtube.com/watch?v=" + link, "Blue", True)
        result_docx.add_paragraph('\n')

        subs_srt = YouTubeTranscriptApi.get_transcript(link, languages=language)
        # variable storing subtitles in srt format

        subs_txt = ''
        # variable storing subtitles in txt format
        for el in subs_srt:
            subs_txt += el['text']
            if (len(subs_txt) > 50000):
                break

        # intro-description of the article
        messages = [  # the instruction for chatGPT without its response
            {"role": "system",
             "content": "Расскажи про что этот текст в 2 предложениях"}
        ]

        messages.append({"role": "user", "content": subs_txt})  

        completion = openai.ChatCompletion.create(  # generation of answer
            model="gpt-3.5-turbo-16k",
            messages=messages,
            n=1,
            temperature=0
        )

        chat_response = completion.choices[0].message.content

        result_docx.add_paragraph(chat_response + '\n\n')

        video_duration = subs_srt[-1]['start']
        # video length

        # number of parts (can/should be edited)
        amount_of_parts = 5  # < 10 minuts - 05 parts
        if video_duration > 600:  # > 10 minuts - 10 parts
            amount_of_parts = 10  # > 01 hour  - 20 parts
        if video_duration > 3600:
            amount_of_parts = 20

        part_duration = math.floor(video_duration / amount_of_parts)
        # length of one part

        subs_txt = ''  # variable storing subtitles in txt format
        num_of_part = 1  
        begining_of_part = int(math.floor(subs_srt[0]['start']))  # start time of the part
        for srt_el in subs_srt:

            subs_txt += srt_el['text']  # collect all the subtitles of the part in subs_txt

            if begining_of_part == -1:  # save the start time of the part
                begining_of_part = int(math.floor(srt_el['start']))

            if srt_el['start'] > part_duration * num_of_part:
                create_part(subs_txt, begining_of_part, num_of_part)  # create a part

                subs_txt = ''
                num_of_part += 1
                begining_of_part = -1
        if subs_txt != '':  # processing of the last "cut" part
            create_part(subs_txt, begining_of_part, num_of_part)

        # save file
        result_file_name = "result_" + link + ".docx"
        result_docx.save(result_file_name)

        # send file to user
        bot.send_document(chat_id=linkk.from_user.id, document=open(result_file_name, "rb"))

    else:
        bot.send_message(linkk.chat.id, "There's something wrong with your link. Try again")


def check_links(links):
    pattern = re.compile(r'^https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+')
    if pattern.match(links):
        return True
    else:
        return False

# bot starting
bot.polling(none_stop=True)
